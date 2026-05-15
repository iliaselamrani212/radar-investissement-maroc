"""
lecture/docx_reader.py - Fonctionnalité 2 : Lecture DOCX
"""
import logging

logger = logging.getLogger(__name__)


def extraire_texte_docx(docx_path: str) -> str:
    """Extrait le texte d'un document Word"""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx non installé. pip install python-docx")
        return ""

    try:
        doc = Document(docx_path)
        paragraphes = [p.text for p in doc.paragraphs if p.text.strip()]

        # Lecture des tableaux aussi
        tableaux = []
        for table in doc.tables:
            for row in table.rows:
                ligne = " | ".join(cell.text.strip() for cell in row.cells)
                if ligne.strip():
                    tableaux.append(ligne)

        return "\n".join(paragraphes + tableaux)
    except Exception as e:
        logger.error(f"Erreur lecture DOCX {docx_path}: {e}")
        return ""
